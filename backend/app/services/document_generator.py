from io import BytesIO
import shutil
import subprocess
import tempfile
from textwrap import wrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONTENT_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}

SECTION_TITLES = {
    "summary": "SUMMARY",
    "skills": "TECHNICAL SKILLS",
    "projects": "PROJECTS",
    "experience": "EXPERIENCE",
    "education": "EDUCATION",
}


def generate_resume_document(
    rendered_content: dict, template_key: str, output_format: str
) -> bytes:
    if output_format == "docx":
        return _generate_docx(rendered_content, template_key)
    if output_format == "pdf":
        if template_key == "ats_classic":
            latex_pdf = _try_generate_latex_pdf(rendered_content)
            if latex_pdf:
                return latex_pdf
        return _generate_pdf(rendered_content, template_key)
    raise ValueError("Unsupported output format")


def get_content_type(output_format: str) -> str:
    return CONTENT_TYPES[output_format]


def _template_settings(template_key: str) -> dict:
    if template_key == "compact":
        return {"font_size": 9.5, "bullet_size": 9.25, "margin": 0.45}
    if template_key == "executive":
        return {"font_size": 10.5, "bullet_size": 10.25, "margin": 0.65}
    return {"font_size": 10.0, "bullet_size": 9.75, "margin": 0.6}


def _section_order(rendered_content: dict) -> list[str]:
    return rendered_content.get("section_order") or [
        "summary",
        "skills",
        "projects",
        "experience",
        "education",
    ]


def _generate_docx(rendered_content: dict, template_key: str) -> bytes:
    settings = _template_settings(template_key)
    document = Document()
    section = document.sections[0]
    margin = Inches(settings["margin"])
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(settings["font_size"])
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(11.5)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 41, 55)

    _add_docx_header(document, rendered_content.get("header") or {}, template_key)

    for section_key in _section_order(rendered_content):
        if section_key == "summary":
            _add_docx_summary(document, rendered_content.get("summary"))
        elif section_key == "skills":
            _add_docx_skills(document, rendered_content.get("skills") or [])
        elif section_key in {"projects", "experience"}:
            _add_docx_entries(
                document,
                SECTION_TITLES[section_key],
                rendered_content.get(section_key) or [],
                settings["bullet_size"],
            )
        elif section_key == "education":
            _add_docx_education(
                document,
                SECTION_TITLES[section_key],
                rendered_content.get("education") or [],
            )

    _add_docx_extra_sections(document, rendered_content.get("extra_sections") or [])
    _add_docx_links(document, rendered_content.get("additional_links") or [])

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _add_docx_header(document: Document, header: dict, template_key: str) -> None:
    name = header.get("name") or "Candidate Name"
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(name)
    run.bold = True
    run.font.size = Pt(18)

    role_bits = [
        value
        for value in [header.get("target_role"), header.get("target_company")]
        if value
    ]
    if role_bits and template_key != "ats_classic":
        role = document.add_paragraph(" at ".join(role_bits))
        role.alignment = WD_ALIGN_PARAGRAPH.CENTER
        role.runs[0].font.size = Pt(10)
        role.runs[0].font.color.rgb = RGBColor(75, 85, 99)

    for links in _header_link_rows(header.get("links") or []):
        link_paragraph = document.add_paragraph()
        link_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        link_paragraph.paragraph_format.space_after = Pt(2)
        for index, link in enumerate(links):
            if index:
                separator = link_paragraph.add_run(" | ")
                separator.font.color.rgb = RGBColor(156, 163, 175)
            _add_docx_link_or_text(link_paragraph, link, 9)


def _add_docx_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(17, 24, 39)


def _add_docx_summary(document: Document, summary: str) -> None:
    if not summary:
        return
    _add_docx_heading(document, SECTION_TITLES["summary"])
    paragraph = document.add_paragraph(summary)
    paragraph.paragraph_format.space_after = Pt(2)


def _add_docx_skills(document: Document, skills: list[dict]) -> None:
    if not skills:
        return
    _add_docx_heading(document, SECTION_TITLES["skills"])
    table = document.add_table(rows=0, cols=2)
    for group in skills:
        row = table.add_row()
        label_cell = row.cells[0]
        value_cell = row.cells[1]
        label_cell.width = Inches(1.1)
        value_cell.width = Inches(5.8)

        label_paragraph = label_cell.paragraphs[0]
        label_paragraph.paragraph_format.space_after = Pt(1)
        label = label_paragraph.add_run(f"{group['label']}:")
        label.bold = True
        label.font.size = Pt(9.75)

        value_paragraph = value_cell.paragraphs[0]
        value_paragraph.paragraph_format.space_after = Pt(1)
        values = value_paragraph.add_run(", ".join(group["items"]))
        values.font.size = Pt(9.75)


def _add_docx_entries(
    document: Document, title: str, entries: list, bullet_size: float
) -> None:
    if not entries:
        return
    _add_docx_heading(document, title)
    for entry in entries:
        title_text = entry.get("title") if isinstance(entry, dict) else ""
        meta = entry.get("meta") if isinstance(entry, dict) else ""
        entry_links = entry.get("links", []) if isinstance(entry, dict) else []
        bullets = entry.get("bullets", []) if isinstance(entry, dict) else [entry]

        if title_text:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(title_text)
            run.bold = True
            run.font.size = Pt(10)
            if meta and isinstance(entry, dict) and entry.get("kind") == "project":
                meta_run = paragraph.add_run(f" ({meta})")
                meta_run.font.size = Pt(9.25)
            if entry_links:
                paragraph.add_run(" | ").font.size = Pt(9)
                _add_docx_inline_links(paragraph, entry_links, 9)

        if meta and not (isinstance(entry, dict) and entry.get("kind") == "project"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(meta)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(75, 85, 99)

        for bullet in bullets:
            text, links = _item_text_and_links(bullet)
            if not text:
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.first_line_indent = Inches(-0.12)
            paragraph.add_run(". ").font.size = Pt(bullet_size)
            paragraph.add_run(text).font.size = Pt(bullet_size)
            if links:
                paragraph.add_run(" (").font.size = Pt(bullet_size)
                _add_docx_inline_links(paragraph, links, bullet_size)
                paragraph.add_run(")").font.size = Pt(bullet_size)


def _add_docx_education(document: Document, title: str, entries: list) -> None:
    if not entries:
        return
    _add_docx_heading(document, title)
    for entry in entries:
        if isinstance(entry, dict) and entry.get("raw"):
            paragraph = document.add_paragraph(entry["raw"])
            paragraph.paragraph_format.space_after = Pt(1)
            continue

        if isinstance(entry, dict):
            heading = " - ".join(
                part for part in [entry.get("degree"), entry.get("institution")] if part
            )
            if heading:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run(heading)
                run.bold = True
                run.font.size = Pt(10)
            if entry.get("meta"):
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(1)
                run = paragraph.add_run(entry["meta"])
                run.italic = True
                run.font.size = Pt(9)
            for detail in entry.get("details") or []:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.left_indent = Inches(0.18)
                paragraph.paragraph_format.first_line_indent = Inches(-0.12)
                paragraph.add_run(". ").font.size = Pt(9.75)
                paragraph.add_run(detail).font.size = Pt(9.75)


def _add_docx_links(document: Document, links: list[dict]) -> None:
    if not links:
        return
    _add_docx_heading(document, "LINKS")
    for link in links:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        _add_docx_link_or_text(paragraph, link, 9.75)


def _add_docx_extra_sections(document: Document, sections: list[dict]) -> None:
    for section in sections:
        title = section.get("title")
        items = section.get("items") or []
        if not title or not items:
            continue
        _add_docx_heading(document, title.upper())
        for item in items:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.first_line_indent = Inches(-0.12)
            paragraph.add_run(". ").font.size = Pt(9.75)
            paragraph.add_run(item).font.size = Pt(9.75)


def _pdf_lines(rendered_content: dict) -> list[tuple[str, str]]:
    header = rendered_content.get("header") or {}
    lines: list[tuple[str, str]] = [
        ("name", header.get("name") or "Candidate Name"),
    ]

    role_bits = [
        value
        for value in [header.get("target_role"), header.get("target_company")]
        if value
    ]
    if role_bits:
        lines.append(("meta", " at ".join(role_bits)))

    links = [link.get("label") for link in header.get("links") or [] if link.get("label")]
    if links:
        lines.append(("meta", " | ".join(links)))

    for section_key in _section_order(rendered_content):
        if section_key == "summary" and rendered_content.get("summary"):
            lines.append(("heading", SECTION_TITLES["summary"]))
            lines.extend(("paragraph", line) for line in wrap(rendered_content["summary"], 96))
        elif section_key == "skills" and rendered_content.get("skills"):
            lines.append(("heading", SECTION_TITLES["skills"]))
            for group in rendered_content["skills"]:
                label = f"{group['label']}:"
                values = ", ".join(group["items"])
                lines.append(("paragraph", f"{label:<14} {values}"))
        elif section_key in {"projects", "experience"} and rendered_content.get(section_key):
            lines.append(("heading", SECTION_TITLES[section_key]))
            for entry in rendered_content[section_key]:
                if isinstance(entry, dict) and entry.get("title"):
                    title = entry["title"]
                    if entry.get("links"):
                        title = f"{title} ({_links_text(entry['links'])})"
                    lines.append(("subheading", title))
                if isinstance(entry, dict) and entry.get("meta"):
                    lines.append(("meta", entry["meta"]))

                bullets = entry.get("bullets", []) if isinstance(entry, dict) else [entry]
                for bullet in bullets:
                    text, links = _item_text_and_links(bullet)
                    if links:
                        text = f"{text} ({_links_text(links)})"
                    wrapped = wrap(text, width=92)
                    if wrapped:
                        lines.append(("bullet", wrapped[0]))
                        lines.extend(("indent", line) for line in wrapped[1:])
        elif section_key == "education" and rendered_content.get("education"):
            lines.append(("heading", SECTION_TITLES["education"]))
            for entry in rendered_content["education"]:
                if isinstance(entry, dict) and entry.get("raw"):
                    lines.append(("paragraph", entry["raw"]))
                    continue
                if isinstance(entry, dict):
                    heading = " - ".join(
                        part
                        for part in [entry.get("degree"), entry.get("institution")]
                        if part
                    )
                    if heading:
                        lines.append(("subheading", heading))
                    if entry.get("meta"):
                        lines.append(("meta", entry["meta"]))
                    for detail in entry.get("details") or []:
                        lines.append(("bullet", detail))

    if rendered_content.get("additional_links"):
        lines.append(("heading", "LINKS"))
        for link in rendered_content["additional_links"]:
            lines.append(("paragraph", link.get("label") or link.get("url") or ""))

    return lines


def _item_text_and_links(item) -> tuple[str, list[dict]]:
    if isinstance(item, dict):
        return item.get("text") or "", item.get("links") or []
    return str(item), []


def _links_text(links: list[dict]) -> str:
    return " | ".join(link.get("label") or link.get("url") or "" for link in links)


def _header_link_rows(links: list[dict]) -> list[list[dict]]:
    contact_types = {"location", "email", "phone"}
    contact_links = [link for link in links if link.get("type") in contact_types]
    profile_links = [link for link in links if link.get("type") not in contact_types]
    return [row for row in [contact_links, profile_links] if row]


def _add_docx_inline_links(paragraph, links: list[dict], size: float) -> None:
    for index, link in enumerate(links):
        if index:
            separator = paragraph.add_run(" | ")
            separator.font.size = Pt(size)
            separator.font.color.rgb = RGBColor(75, 85, 99)
        _add_docx_link_or_text(paragraph, link, size)


def _add_docx_link_or_text(paragraph, link: dict, size: float):
    label = link.get("label") or link.get("url") or ""
    url = link.get("url") or ""
    if url:
        return _add_docx_hyperlink(paragraph, label, url, size)

    run = paragraph.add_run(label)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(55, 65, 81)
    return run


def _add_docx_hyperlink(paragraph, text: str, url: str, size: float):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F2937")
    properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)

    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(int(size * 2)))
    properties.append(font_size)

    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _escape_pdf_uri(uri: str) -> str:
    return _escape_pdf_text(uri)


def _pdf_text_width(text: str, size: float) -> float:
    return len(text) * size * 0.55


def _try_generate_latex_pdf(rendered_content: dict) -> bytes | None:
    engine = shutil.which("pdflatex")
    if not engine:
        return None

    tex = _latex_document(rendered_content)
    with tempfile.TemporaryDirectory() as temp_dir:
        tex_path = f"{temp_dir}/resume.tex"
        pdf_path = f"{temp_dir}/resume.pdf"
        with open(tex_path, "w", encoding="utf-8") as handle:
            handle.write(tex)

        result = subprocess.run(
            [
                engine,
                "-interaction=nonstopmode",
                "-halt-on-error",
                "-output-directory",
                temp_dir,
                tex_path,
            ],
            capture_output=True,
            timeout=20,
            check=False,
        )
        if result.returncode != 0:
            return None

        with open(pdf_path, "rb") as handle:
            return handle.read()


def _latex_document(rendered_content: dict) -> str:
    header = rendered_content.get("header") or {}
    links = header.get("links") or []
    contacts = _links_by_type(links)

    sections = [
        _latex_summary(rendered_content.get("summary")),
        _latex_skills(rendered_content.get("skills") or []),
        _latex_entries("Projects", rendered_content.get("projects") or []),
        _latex_entries("Experience", rendered_content.get("experience") or []),
        _latex_education(rendered_content.get("education") or []),
    ]
    for section in rendered_content.get("extra_sections") or []:
        sections.append(_latex_item_section(section.get("title"), section.get("items") or []))
    sections.append(_latex_link_section(rendered_content.get("additional_links") or []))

    return "\n".join(
        [
            r"\documentclass[a4paper,10pt]{article}",
            r"\usepackage[empty]{fullpage}",
            r"\usepackage{titlesec}",
            r"\usepackage{enumitem}",
            r"\usepackage{parskip}",
            r"\usepackage{fontawesome5}",
            r"\usepackage[colorlinks=true, urlcolor=blue]{hyperref}",
            r"\usepackage{tabularx}",
            r"\titleformat{\section}{\large\bfseries}{}{0em}{}",
            r"\setlist[itemize]{noitemsep, topsep=0pt}",
            r"\begin{document}",
            r"\begin{center}",
            rf"{{\LARGE \textbf{{{_latex_escape(header.get('name') or 'Candidate Name')}}}}}\\[15pt]",
            r"\begin{tabular}{l l l}",
            _latex_header_contact_row(contacts),
            r"\\[12pt]",
            _latex_header_profile_row(contacts),
            r"\end{tabular}",
            r"\end{center}",
            *[section for section in sections if section],
            r"\end{document}",
        ]
    )


def _links_by_type(links: list[dict]) -> dict[str, dict]:
    return {link.get("type", ""): link for link in links}


def _latex_header_contact_row(links: dict[str, dict]) -> str:
    location = links.get("location") or {"label": ""}
    email = links.get("email") or {"label": "", "url": ""}
    phone = links.get("phone") or {"label": "", "url": ""}
    return " & ".join(
        [
            rf"\faIcon{{map-marker-alt}}~{_latex_escape(location.get('label') or '')} \hspace{{3cm}}",
            rf"\faIcon{{envelope}}~{_latex_link(email)} \hspace{{2cm}}",
            rf"\faIcon{{phone}}~{_latex_link(phone)}",
        ]
    )


def _latex_header_profile_row(links: dict[str, dict]) -> str:
    linkedin = links.get("linkedin") or {"label": "LinkedIn", "url": ""}
    github = links.get("github") or {"label": "GitHub", "url": ""}
    portfolio = links.get("portfolio") or {"label": "Portfolio", "url": ""}
    return " & ".join(
        [
            rf"\faIcon{{linkedin}}~{_latex_link(linkedin)} \hspace{{2cm}}",
            rf"\faIcon{{github}}~{_latex_link(github)} \hspace{{2cm}}",
            rf"\faIcon{{globe}}~{_latex_link(portfolio)}",
        ]
    )


def _latex_link(link: dict) -> str:
    label = _latex_escape(link.get("label") or link.get("url") or "")
    url = link.get("url") or ""
    if not url:
        return label
    return rf"\href{{{_latex_escape_url(url)}}}{{{label}}}"


def _latex_summary(summary: str | None) -> str:
    if not summary:
        return ""
    return "\n".join([r"\section*{Summary}", _latex_escape(summary)])


def _latex_skills(skills: list[dict]) -> str:
    if not skills:
        return ""
    rows = [
        rf"\textbf{{{_latex_escape(group.get('label') or '')}:}} & {_latex_escape(', '.join(group.get('items') or []))} \\"
        for group in skills
    ]
    return "\n".join(
        [r"\section*{Technical Skills}", r"\begin{tabular}{@{} l l @{}}", *rows, r"\end{tabular}"]
    )


def _latex_entries(title: str, entries: list[dict]) -> str:
    if not entries:
        return ""
    parts = [rf"\section*{{{_latex_escape(title)}}}"]
    for entry in entries:
        entry_title = entry.get("title") or ""
        meta = entry.get("meta") or ""
        left = rf"\textbf{{{_latex_escape(entry_title)}}}"
        if meta and entry.get("kind") == "project":
            left += rf" ({_latex_escape(meta)})"
        right = _latex_links_text(entry.get("links") or [])
        parts.extend(
            [
                r"\begin{tabular*}{\textwidth}{l@{\extracolsep{\fill}}r}",
                rf"{left} & {right} \\",
                r"\end{tabular*}",
            ]
        )
        if meta and entry.get("kind") != "project":
            parts.append(_latex_escape(meta))
        bullets = [(_item_text_and_links(item)[0] or "") for item in entry.get("bullets") or []]
        if bullets:
            parts.append(r"\begin{itemize}")
            parts.extend(rf"\item {_latex_escape(bullet)}" for bullet in bullets if bullet)
            parts.append(r"\end{itemize}")
    return "\n".join(parts)


def _latex_links_text(links: list[dict]) -> str:
    return " ".join(rf"\faIcon{{github}} {_latex_link(link)}" for link in links)


def _latex_education(entries: list[dict]) -> str:
    if not entries:
        return ""
    parts = [r"\section*{Education}"]
    for entry in entries:
        if entry.get("raw"):
            parts.append(_latex_escape(entry["raw"]))
            continue
        heading = " - ".join(
            part for part in [entry.get("degree"), entry.get("institution")] if part
        )
        if heading:
            parts.append(rf"\textbf{{{_latex_escape(heading)}}}")
        if entry.get("meta"):
            parts.append(_latex_escape(entry["meta"]))
        if entry.get("details"):
            parts.append(r"\begin{itemize}")
            parts.extend(rf"\item {_latex_escape(detail)}" for detail in entry["details"])
            parts.append(r"\end{itemize}")
    return "\n".join(parts)


def _latex_item_section(title: str | None, items: list[str]) -> str:
    if not title or not items:
        return ""
    parts = [rf"\section*{{{_latex_escape(title)}}}", r"\begin{itemize}"]
    parts.extend(rf"\item {_latex_escape(item)}" for item in items)
    parts.append(r"\end{itemize}")
    return "\n".join(parts)


def _latex_link_section(links: list[dict]) -> str:
    if not links:
        return ""
    parts = [r"\section*{Links}", r"\begin{itemize}"]
    parts.extend(rf"\item {_latex_link(link)}" for link in links)
    parts.append(r"\end{itemize}")
    return "\n".join(parts)


def _latex_escape(value: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in str(value or ""))


def _latex_escape_url(url: str) -> str:
    return str(url or "").replace("\\", "")


def _generate_pdf(rendered_content: dict, template_key: str) -> bytes:
    objects: list[bytes] = []
    page_refs = []
    page_streams: list[tuple[bytes, list[dict]]] = []
    header = rendered_content.get("header") or {}
    y = 760
    commands: list[str] = []
    annotations: list[dict] = []

    def add_object(body: str | bytes) -> int:
        if isinstance(body, str):
            body = body.encode("latin-1", errors="replace")
        objects.append(body)
        return len(objects)

    def flush_page() -> None:
        nonlocal commands, y, annotations
        if commands:
            page_streams.append(
                (
                    "\n".join(commands).encode("latin-1", errors="replace"),
                    annotations,
                )
            )
        commands = []
        annotations = []
        y = 760

    def ensure_space(height: int) -> None:
        if y - height < 48:
            flush_page()

    def draw_text(text: str, x: float, size: float = 9.5, font: str = "F1") -> float:
        nonlocal y
        commands.append(
            f"BT /{font} {size} Tf {x:.2f} {y:.2f} Td ({_escape_pdf_text(text)}) Tj ET"
        )
        return _pdf_text_width(text, size)

    def draw_center(text: str, size: float = 10, font: str = "F1") -> None:
        nonlocal y
        approx_width = _pdf_text_width(text, size)
        x = max(50, (612 - approx_width) / 2)
        draw_text(text, x, size, font)

    def draw_link_text(
        text: str, url: str, x: float, size: float = 9.5, font: str = "F1"
    ) -> float:
        width = draw_text(text, x, size, font)
        if url:
            annotations.append(
                {
                    "url": url,
                    "rect": [
                        round(x, 2),
                        round(y - 2, 2),
                        round(x + width, 2),
                        round(y + size + 2, 2),
                    ],
                }
            )
        return width

    def draw_link_row(links: list[dict], size: float = 8.5) -> None:
        labels = [link.get("label") or link.get("url") or "" for link in links]
        if not labels:
            return

        separator = " | "
        total_width = sum(_pdf_text_width(label, size) for label in labels)
        total_width += _pdf_text_width(separator, size) * (len(labels) - 1)
        x = max(50, (612 - total_width) / 2)

        for index, link in enumerate(links):
            if index:
                x += draw_text(separator, x, size, "F1")
            x += draw_link_text(
                link.get("label") or link.get("url") or "",
                link.get("url") or "",
                x,
                size,
                "F1",
            )

    def next_line(amount: int = 13) -> None:
        nonlocal y
        y -= amount

    def section_heading(title: str) -> None:
        nonlocal y
        ensure_space(32)
        next_line(10)
        draw_text(title, 50, 10, "F2")
        next_line(14)

    def paragraph(text: str) -> None:
        for line in wrap(text, width=102):
            ensure_space(14)
            draw_text(line, 50, 9.5)
            next_line(13)

    def bullet(text: str) -> None:
        wrapped = wrap(text, width=95)
        for index, line in enumerate(wrapped):
            ensure_space(14)
            prefix = ". " if index == 0 else "  "
            draw_text(f"{prefix}{line}", 62, 9.25)
            next_line(13)

    name = header.get("name") or "Candidate Name"
    draw_center(name, 18, "F2")
    next_line(17)

    role_bits = [
        value
        for value in [header.get("target_role"), header.get("target_company")]
        if value
    ]
    if role_bits and template_key != "ats_classic":
        draw_center(" at ".join(role_bits), 10, "F1")
        next_line(14)

    for links in _header_link_rows(header.get("links") or []):
        draw_link_row(links, 8.5)
        next_line(14)

    next_line(4)

    for section_key in _section_order(rendered_content):
        if section_key == "summary" and rendered_content.get("summary"):
            section_heading(SECTION_TITLES["summary"])
            paragraph(rendered_content["summary"])
        elif section_key == "skills" and rendered_content.get("skills"):
            section_heading(SECTION_TITLES["skills"])
            for group in rendered_content["skills"]:
                ensure_space(14)
                draw_text(f"{group['label']}:", 50, 9.25, "F2")
                draw_text(", ".join(group["items"]), 132, 9.25, "F1")
                next_line(13)
        elif section_key in {"projects", "experience"} and rendered_content.get(section_key):
            section_heading(SECTION_TITLES[section_key])
            for entry in rendered_content[section_key]:
                ensure_space(34)
                if isinstance(entry, dict) and entry.get("title"):
                    x = 50
                    title = entry["title"]
                    if entry.get("meta") and entry.get("kind") == "project":
                        title = f"{title} ({entry['meta']})"
                    x += draw_text(title, x, 9.75, "F2")
                    if entry.get("links"):
                        link_labels = [
                            link.get("label") or link.get("url") or ""
                            for link in entry["links"]
                        ]
                        link_text = " | ".join(link_labels)
                        link_x = max(x + 12, 562 - _pdf_text_width(link_text, 9.25))
                        for link_index, link in enumerate(entry["links"]):
                            if link_index:
                                link_x += draw_text(" | ", link_x, 9.25, "F1")
                            link_x += draw_link_text(
                                link.get("label") or link.get("url") or "",
                                link.get("url") or "",
                                link_x,
                                9.25,
                                "F1",
                            )
                    next_line(12)
                if (
                    isinstance(entry, dict)
                    and entry.get("meta")
                    and entry.get("kind") != "project"
                ):
                    draw_text(entry["meta"], 50, 8.75, "F3")
                    next_line(12)

                bullets = entry.get("bullets", []) if isinstance(entry, dict) else [entry]
                for item in bullets:
                    text, item_links = _item_text_and_links(item)
                    if item_links:
                        text = f"{text} ({_links_text(item_links)})"
                    if text:
                        bullet(text)
                next_line(3)
        elif section_key == "education" and rendered_content.get("education"):
            section_heading(SECTION_TITLES["education"])
            for entry in rendered_content["education"]:
                if isinstance(entry, dict) and entry.get("raw"):
                    paragraph(entry["raw"])
                    continue
                if isinstance(entry, dict):
                    heading = " - ".join(
                        part
                        for part in [entry.get("degree"), entry.get("institution")]
                        if part
                    )
                    if heading:
                        draw_text(heading, 50, 9.5, "F2")
                        next_line(12)
                    if entry.get("meta"):
                        draw_text(entry["meta"], 50, 8.75, "F3")
                        next_line(12)
                    for detail in entry.get("details") or []:
                        bullet(detail)

    for section in rendered_content.get("extra_sections") or []:
        title = section.get("title")
        items = section.get("items") or []
        if not title or not items:
            continue
        section_heading(title.upper())
        for item in items:
            bullet(item)

    if rendered_content.get("additional_links"):
        section_heading("LINKS")
        for link in rendered_content["additional_links"]:
            ensure_space(14)
            draw_link_text(
                link.get("label") or link.get("url") or "",
                link.get("url") or "",
                50,
                9.5,
                "F1",
            )
            next_line(13)

    flush_page()

    catalog_id = add_object("<< /Type /Catalog /Pages 2 0 R >>")
    pages_id = add_object(b"")
    font_id = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    bold_font_id = add_object(
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>"
    )
    italic_font_id = add_object(
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Oblique >>"
    )

    for stream, page_annotations in page_streams:
        content_id = add_object(
            b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        )
        annotation_refs = []
        for annotation in page_annotations:
            rect = " ".join(str(value) for value in annotation["rect"])
            annotation_refs.append(
                add_object(
                    "<< /Type /Annot /Subtype /Link "
                    f"/Rect [{rect}] /Border [0 0 0] "
                    f"/A << /S /URI /URI ({_escape_pdf_uri(annotation['url'])}) >> >>"
                )
            )
        annots = (
            f"/Annots [{' '.join(f'{ref} 0 R' for ref in annotation_refs)}] "
            if annotation_refs
            else ""
        )
        page_id = add_object(
            f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 612 792] "
            f"{annots}"
            f"/Resources << /Font << /F1 {font_id} 0 R /F2 {bold_font_id} 0 R /F3 {italic_font_id} 0 R >> >> "
            f"/Contents {content_id} 0 R >>"
        )
        page_refs.append(page_id)

    kids = " ".join(f"{page_id} 0 R" for page_id in page_refs)
    objects[pages_id - 1] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_refs)} >>".encode(
        "latin-1"
    )

    output = BytesIO()
    output.write(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(output.tell())
        output.write(f"{index} 0 obj\n".encode("ascii"))
        output.write(body)
        output.write(b"\nendobj\n")

    xref_offset = output.tell()
    output.write(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.write(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.write(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_id} 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF"
        ).encode("ascii")
    )
    return output.getvalue()
