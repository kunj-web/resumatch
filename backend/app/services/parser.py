import pdfplumber
import io
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    parts = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                parts.append(page_text)

            link_lines = _extract_pdf_link_lines(page)
            if link_lines:
                parts.extend(link_lines)

    return "\n".join(parts).strip()


def _extract_pdf_link_lines(page) -> list[str]:
    link_lines = []
    hyperlinks = getattr(page, "hyperlinks", None) or []

    for link in hyperlinks:
        uri = link.get("uri")
        if not uri:
            continue

        label = _extract_link_label(page, link)
        if label:
            link_lines.append(f"{label}: {uri}")
        else:
            link_lines.append(uri)

    return list(dict.fromkeys(link_lines))


def _extract_link_label(page, link: dict) -> str:
    words = page.extract_words() or []
    top = link.get("top")
    bottom = link.get("bottom")
    if top is None or bottom is None:
        return ""

    same_line_words = [
        word
        for word in words
        if word.get("top", 0) <= bottom + 2 and word.get("bottom", 0) >= top - 2
    ]
    if same_line_words:
        return " ".join(word["text"] for word in same_line_words).strip()

    link_words = [
        word
        for word in words
        if word.get("x0", 0) <= link.get("x1", 0) + 2
        and word.get("x1", 0) >= link.get("x0", 0) - 2
        and word.get("top", 0) <= bottom + 2
        and word.get("bottom", 0) >= top - 2
    ]
    return " ".join(word["text"] for word in link_words).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    parts = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def clean_text(text: str) -> str:
    # Remove excessive whitespace and blank lines
    lines = [line.strip() for line in text.splitlines()]
    non_empty = [line for line in lines if line]
    return "\n".join(non_empty)


def parse_resume(file_bytes: bytes, file_type: str) -> str:
    if file_type == "pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    elif file_type == "docx":
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported resume file type")

    cleaned = clean_text(raw_text)
    return cleaned
